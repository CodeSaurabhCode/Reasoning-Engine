
from docx import Document
from docx.shared import Pt, RGBColor
import ast
import io

def create_docx(data) -> bytes:
    doc = Document()
    for key, values in data.items():
        heading = doc.add_paragraph()
        heading_run = heading.add_run(f"<{key}>")
        heading_run.font.size = Pt(13)
        heading_run.font.color.rgb = RGBColor(150, 150, 150)

        if isinstance(values, dict):
            for k, v in values.items():
                para = doc.add_paragraph()
                run = para.add_run(str(v))
                run.font.size = Pt(12) 

        elif isinstance(values, str):
            try:
                string_dict = ast.literal_eval(values)
                if isinstance(string_dict, dict):
                    max_len = max(len(string_dict.get("headline", [])), len(string_dict.get("subheadline", [])))
                    for i in range(max_len):
                        if "headline" in string_dict and i < len(string_dict["headline"]):
                            headline_para = doc.add_paragraph()
                            headline_run = headline_para.add_run(string_dict["headline"][i])
                            headline_run.bold = True
                            headline_run.font.size = Pt(13)
                        if "subheadline" in string_dict and i < len(string_dict["subheadline"]):
                            subheadline_para = doc.add_paragraph()
                            subheadline_run = subheadline_para.add_run(string_dict["subheadline"][i])
                            subheadline_run.font.size = Pt(12)
            except (ValueError, SyntaxError):
                para = doc.add_paragraph()
                run = para.add_run(str(values))
                run.font.size = Pt(12)
        else:
            para = doc.add_paragraph()
            run = para.add_run(str(values))
            run.font.size = Pt(12)
        doc.add_paragraph("\n")
    
    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
