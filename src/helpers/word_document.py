import io
from docx import Document
from docx.shared import Pt, RGBColor
import ast
def create_docx(data, filename="web_copy.docx"):
    doc = Document()
    for key, values in data.items():
        heading = doc.add_paragraph()
        heading_run = heading.add_run(f"<{key}>")
        heading_run.font.size = Pt(13)
        heading_run.font.color.rgb = RGBColor(150, 150, 150)

        if isinstance(values, dict):
            for key, value in values:
                para = doc.add_paragraph()
                run = para.add_run(str(value))
                run.font.size = Pt(12) 
        if isinstance(values, str):
            try:
                string_dict = ast.literal_eval(values)
                if isinstance(string_dict, dict):
                    # Handle headline
                    max_len = max(len(string_dict.get("headline", [])), len(string_dict.get("subheadline", [])))

                    # Loop through indices and alternate between headline and subheadline
                    for i in range(max_len):
                        # Add headline if available
                        if "headline" in string_dict and i < len(string_dict["headline"]):
                            headline_para = doc.add_paragraph()
                            headline_run = headline_para.add_run(string_dict["headline"][i])
                            headline_run.bold = True
                            headline_run.font.size = Pt(13)

                        # Add subheadline if available
                        if "subheadline" in string_dict and i < len(string_dict["subheadline"]):
                            subheadline_para = doc.add_paragraph()
                            subheadline_run = subheadline_para.add_run(string_dict["subheadline"][i])
                            subheadline_run.font.size = Pt(12)

            except (ValueError, SyntaxError):
                # If it's not a valid string-dict, treat as plain content
                para = doc.add_paragraph()
                run = para.add_run(str(values))
                run.font.size = Pt(12)
        else:
            para = doc.add_paragraph()
            run = para.add_run(str(values))
            run.font.size = Pt(12)
        doc.add_paragraph("\n")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()